from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

document = Document()
document.styles['Normal'].font.name = u'宋体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['Normal'].font.size = Pt(10.5)
document.styles['Normal'].font.color.rgb = RGBColor(255, 0, 0)

document.styles['Body Text'].font.name = u'宋体'
document.styles['Body Text']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['Body Text'].font.size = Pt(10.5)
document.styles['Body Text'].font.color.rgb = RGBColor(100, 0, 0)
paragraph = document.add_paragraph('在一')
paragraph = document.add_paragraph('在一', 'Body Text')
document.save('demo.docx')