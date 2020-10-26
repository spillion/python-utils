from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
import re
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

file1 = open("C:\\Users\\kedacom\\Desktop\\11.txt", "r", encoding='utf-8')
answers = file1.read().replace('\n', '').replace('\r', '')
print(answers)
file1.close()

file2 = open("C:\\Users\\kedacom\\Desktop\\12.txt", "r", encoding='utf-8')
titles = file2.read().replace('\n', '').replace('\r', '')
print(titles)
file2.close()

title = re.split(r'[0-9]{1,2}\.', titles)
for item in title:
    print('111' + item)

answer = re.split(r'[0-9]{1,2}\. 解答：', answers)
for item in answer:
    print('222' + item)

print(len(title))
print(len(answer))

document = Document()
style = document.styles
print(style)

document.styles['Normal'].font.name = u'宋体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['Normal'].font.size = Pt(10.5)
document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
document.styles['Normal'].paragraph_format.space_after = Pt(0)

document.styles['Body Text'].font.name = u'宋体'
document.styles['Body Text']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['Body Text'].font.size = Pt(18)
document.styles['Body Text'].font.color.rgb = RGBColor(255, 0, 0)

document.styles['No Spacing'].font.name = u'宋体'
document.styles['No Spacing']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['No Spacing'].font.size = Pt(12)
document.styles['No Spacing'].font.color.rgb = RGBColor(255, 0, 0)

if len(title) == len(answer):
    for item in range(0, len(title)):
        document.add_paragraph('【填空题】（每空10分）', 'Body Text')
        document.add_paragraph(title[item])
        document.add_paragraph('')
        document.add_paragraph('判分时答案内字母区分大小写：否', 'No Spacing')
        document.add_paragraph('判分时区分多个空的先后顺序：否', 'No Spacing')
        document.add_paragraph('所属章：5.2', 'No Spacing')
        document.add_paragraph('难度：一般', 'No Spacing')
        paragraph = document.add_paragraph()
        run = paragraph.add_run('解析：')
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(255, 0, 0)
        paragraph.add_run(answer[item])
        paragraph = document.add_paragraph()
        run = paragraph.add_run('知识点：')
        run.font.color.rgb = RGBColor(255, 0, 0)
        document.add_paragraph('')
else:
    print('error')
    print(len(title))
    print(len(answer))

document.save('6.5.docx')
